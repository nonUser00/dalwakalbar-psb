import docx
import pandas as pd
import glob
import os

folder = 'kebutuhan client'
docx_files = glob.glob(f'{folder}/*.docx')
xlsx_files = glob.glob(f'{folder}/*.xlsx')

print("=== DOCX FILES ===")
for f in docx_files:
    print(f"\n--- {f} ---")
    try:
        doc = docx.Document(f)
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                print(p.text.strip())
        
        for table in doc.tables:
            print("[TABLE]")
            for row in table.rows:
                print(" | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells]))
    except Exception as e:
        print(f"Error reading {f}: {e}")

print("\n=== XLSX FILES ===")
for f in xlsx_files:
    print(f"\n--- {f} ---")
    try:
        xls = pd.ExcelFile(f)
        for sheet_name in xls.sheet_names:
            print(f"\nSheet: {sheet_name}")
            df = pd.read_excel(xls, sheet_name=sheet_name)
            print(df.head(10).to_string())
    except Exception as e:
        print(f"Error reading {f}: {e}")
